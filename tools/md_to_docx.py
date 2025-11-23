from pathlib import Path
from markdown import markdown
from docx import Document
from bs4 import BeautifulSoup

md_file = Path('PROJECT_PROPOSAL_BMW_SALES.md')
docx_file = Path('PROJECT_PROPOSAL_BMW_SALES.docx')

if not md_file.exists():
    raise SystemExit(f'Markdown file not found: {md_file}')

md_text = md_file.read_text(encoding='utf-8')

# Render markdown to HTML
html = markdown(md_text, extensions=['extra', 'sane_lists'])

# Parse HTML
soup = BeautifulSoup(html, 'html.parser')

doc = Document()

def add_html_to_doc(element, container):
    """Recursively add HTML element content to the DOCX document."""
    for child in element.children:
        if child.name is None:
            # plain text
            text = child.string.strip() if child.string else ''
            if text:
                container.add_paragraph(text)
        elif child.name in ['h1', 'h2', 'h3']:
            level = 1 if child.name == 'h1' else (2 if child.name == 'h2' else 3)
            container.add_heading(child.get_text().strip(), level=level)
        elif child.name == 'p':
            container.add_paragraph(child.get_text().strip())
        elif child.name in ['ul', 'ol']:
            for li in child.find_all('li', recursive=False):
                style = 'List Number' if child.name == 'ol' else 'List Bullet'
                container.add_paragraph(li.get_text().strip(), style=style)
        elif child.name == 'strong' or child.name == 'b':
            p = container.add_paragraph()
            p.add_run(child.get_text().strip()).bold = True
        elif child.name == 'em' or child.name == 'i':
            p = container.add_paragraph()
            p.add_run(child.get_text().strip()).italic = True
        else:
            # For any other tag, recurse
            add_html_to_doc(child, container)

add_html_to_doc(soup, doc)

doc.save(docx_file)
print(f'Wrote: {docx_file.resolve()}')
